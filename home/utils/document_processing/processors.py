"""
Document processors for handling different file types.
"""
import os
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional

import PyPDF2
from docx import Document as DocxDocument
import mammoth
from bs4 import BeautifulSoup
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Initialize NLTK resources
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

logger = logging.getLogger(__name__)

class BaseDocumentProcessor(ABC):
    """Base class for document processors"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract raw text from document"""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document"""
        pass
    
    def process(self, file_path: str) -> Dict[str, Any]:
        """Process document and return structured content"""
        try:
            text = self.extract_text(file_path)
            metadata = self.extract_metadata(file_path)
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Basic sentence tokenization
            sentences = sent_tokenize(text)
            
            # Return structured result
            return {
                'metadata': metadata,
                'text': text,
                'chunks': chunks,
                'sentences': sentences,
                'word_count': len(word_tokenize(text))
            }
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

class PDFProcessor(BaseDocumentProcessor):
    """Processor for PDF documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF document"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
            raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF document"""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                info = reader.metadata
                
                # Convert PyPDF2 metadata to dict
                metadata = {}
                if info:
                    for key in info:
                        metadata[key] = info[key]
                
                # Add page count
                metadata['page_count'] = len(reader.pages)
                
                return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata from PDF {file_path}: {str(e)}")
            raise
    
    def extract_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract sections from PDF document (e.g., headers, paragraphs)"""
        # This is a simplified implementation - real-world implementation would
        # need more sophisticated logic to identify sections in PDFs
        sections = []
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    
                    # Simple heuristic - split by double newlines
                    parts = text.split('\n\n')
                    for part in parts:
                        if part.strip():
                            sections.append({
                                'type': 'paragraph',
                                'content': part.strip(),
                                'page': page_num + 1
                            })
            
            return sections
        except Exception as e:
            logger.error(f"Error extracting sections from PDF {file_path}: {str(e)}")
            raise

class DocxProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX document"""
        try:
            doc = DocxDocument(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            # Fall back to mammoth if python-docx fails
            try:
                with open(file_path, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html = result.value
                    soup = BeautifulSoup(html, 'html.parser')
                    return soup.get_text()
            except Exception as mammoth_err:
                logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}, Mammoth error: {str(mammoth_err)}")
                raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract core properties
            core_props = {}
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                attributes = ['author', 'category', 'comments', 'content_status',
                             'created', 'identifier', 'keywords', 'language',
                             'last_modified_by', 'last_printed', 'modified',
                             'revision', 'subject', 'title', 'version']
                
                for attr in attributes:
                    if hasattr(props, attr):
                        prop_value = getattr(props, attr)
                        if prop_value is not None:
                            core_props[attr] = str(prop_value)
            
            # Add document statistics
            stats = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            return {**core_props, **stats}
        except Exception as e:
            logger.error(f"Error extracting metadata from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_sections(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract structured sections from DOCX document"""
        sections = []
        
        try:
            doc = DocxDocument(file_path)
            
            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    section_type = 'paragraph'
                    
                    # Check if it's a heading
                    if para.style.name.startswith('Heading'):
                        section_type = para.style.name
                    
                    sections.append({
                        'type': section_type,
                        'content': para.text,
                        'index': i
                    })
            
            # Process tables (simplified)
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                sections.append({
                    'type': 'table',
                    'content': table_data,
                    'index': i
                })
            
            return sections
        except Exception as e:
            logger.error(f"Error extracting sections from DOCX {file_path}: {str(e)}")
            raise

class TxtProcessor(BaseDocumentProcessor):
    """Processor for plain text documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from TXT document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
                raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from TXT document"""
        try:
            stats = os.stat(file_path)
            return {
                'file_size': stats.st_size,
                'created': stats.st_ctime,
                'modified': stats.st_mtime
            }
        except Exception as e:
            logger.error(f"Error extracting metadata from TXT {file_path}: {str(e)}")
            raise

class HtmlProcessor(BaseDocumentProcessor):
    """Processor for HTML documents"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from HTML document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html = file.read()
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        except UnicodeDecodeError:
            # Try with different encodings if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    html = file.read()
                    soup = BeautifulSoup(html, 'html.parser')
                    return soup.get_text()
            except Exception as e:
                logger.error(f"Error extracting text from HTML {file_path}: {str(e)}")
                raise
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from HTML document"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html = file.read()
                soup = BeautifulSoup(html, 'html.parser')
                
                metadata = {
                    'file_size': os.path.getsize(file_path),
                    'created': os.path.getctime(file_path),
                    'modified': os.path.getmtime(file_path)
                }
                
                # Extract meta tags
                meta_tags = {}
                for meta in soup.find_all('meta'):
                    name = meta.get('name') or meta.get('property')
                    content = meta.get('content')
                    if name and content:
                        meta_tags[name] = content
                
                # Get title
                title_tag = soup.find('title')
                if title_tag:
                    metadata['title'] = title_tag.string
                
                return {**metadata, 'meta_tags': meta_tags}
        except Exception as e:
            logger.error(f"Error extracting metadata from HTML {file_path}: {str(e)}")
            raise

def get_processor(file_type: str) -> BaseDocumentProcessor:
    """Factory function to get the appropriate processor for a file type"""
    processors = {
        'pdf': PDFProcessor,
        'docx': DocxProcessor,
        'txt': TxtProcessor,
        'html': HtmlProcessor,
    }
    
    processor_class = processors.get(file_type.lower())
    if not processor_class:
        raise ValueError(f"Unsupported file type: {file_type}")
    
    return processor_class()